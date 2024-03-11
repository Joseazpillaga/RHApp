import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import subprocess
import sys
import os
import re

def instalar_paquete(paquete):
    subprocess.check_call([sys.executable, "-m", "pip", "install", paquete])

def verificar_instalacion(paquete):
    try:
        importlib.metadata.version(paquete)
        return True
    except importlib.metadata.PackageNotFoundError:
        return False

def verificar_y_instalar(paquete):
    if not verificar_instalacion(paquete):
        print(f"Instalando {paquete}...")
        instalar_paquete(paquete)
        print(f"{paquete} instalado correctamente.")

def llenar_formulario(documento, datos_usuario):
    # Reemplazar los marcadores en el documento con la información proporcionada
    for p in documento.paragraphs:
        for marcador, entry in datos_usuario.items():
            valor = entry.get()  # Obtener el texto introducido por el usuario
            p.text = re.sub(re.escape(marcador), valor, p.text)
    # No es necesario guardar el documento aquí, ya que se espera que el documento sea guardado fuera de esta función

def procesar_archivos(archivos, datos_usuario):
    # Obtener la ruta del directorio del script ejecutable
    directorio_script = os.path.dirname(sys.argv[0])
    # Construir la ruta para la carpeta "Procesados"
    carpeta_procesados = os.path.join(directorio_script, "Procesados")
    # Verificar si la carpeta "Procesados" existe, si no, crearla
    if not os.path.exists(carpeta_procesados):
        os.makedirs(carpeta_procesados)

    for archivo in archivos:
        print(f"Procesando archivo: {archivo}")
        if not os.path.exists(archivo):
            print(f"El archivo {archivo} no existe.")
            continue

        document = Document(archivo)
        archivo_modificado = os.path.join(carpeta_procesados, os.path.basename(archivo).replace('.docx', '_modificado.docx'))  # Nombre del archivo modificado en la carpeta "Procesados"

        # Llenar el formulario en el documento antes de guardar
        llenar_formulario(document, datos_usuario)

        # Verificar si los marcadores se están reemplazando correctamente
        print("Marcadores reemplazados en el documento:")
        for p in document.paragraphs:
            print(p.text)

        # Guardar el documento con los marcadores reemplazados
        document.save(archivo_modificado)
        print(f"Documento guardado como: {archivo_modificado}")  # Verificar la ruta del archivo modificado
        messagebox.showinfo("Documento actualizado", f"Documento '{os.path.basename(archivo)}' actualizado y guardado en la carpeta 'Procesados'.")

    print("Todos los archivos procesados.")

def seleccionar_archivos():
    archivos = filedialog.askopenfilenames(title="Seleccionar archivos", filetypes=(("Documentos Word", "*.docx"), ("Todos los archivos", "*.*")))
    if archivos:
        procesar_archivos(archivos, datos_usuario)

def borrar_campos():
    for entry in datos_usuario.values():
        entry.delete(0, tk.END)

def salir():
    root.destroy()

def abrir_carpeta_procesados():
    # Obtener la ruta del directorio del script ejecutable
    directorio_script = os.path.dirname(sys.argv[0])
    # Construir la ruta para la carpeta "Procesados"
    carpeta_procesados = os.path.join(directorio_script, "Procesados")
    # Abrir la carpeta en el explorador de archivos del sistema operativo
    subprocess.Popen(f'explorer "{carpeta_procesados}"')

# Crear la ventana principal
root = tk.Tk()
root.iconbitmap("C:\\Users\\joseelias.azpillaga\\OneDrive - METRO\\Escritorio\\APP RRHH\\RRHHAPP\\icon.ico")
root.title("Altas P&C")
root.geometry("400x500") #Tamaño de ventana
root.resizable(False, False)

# Definir nombres personalizados para las variables para la interfaz gráfica
nombres_personalizados = [
    "Nombre",
    "Fecha de hoy",
    "NIF",
    "Fecha de nacimiento",
    "Numero Seguridad Social",
    "Nivel de educacion",
    "Nacionalidad",
    "Domicilio",
    "Pais de residencia",
    "Contratado como",
    "Grupo profesional",
    "Funciones",
    "Fecha de inicio",
    "Periodo de prueba",
    "Salario"
]

# Definir nombres de variables internas, las que deben tener los archivos docx
nombres_variables = [
    "Name",
    "TodayDate",
    "NIF",
    "BornDate",
    "SsNum",
    "EducationLevel",
    "Nationality",
    "Address",
    "Country",
    "HiredAs",
    "ProfGroup",
    "Functions",
    "StartDate",
    "Probation",
    "Salary"
]

# Solicitar información al usuario
datos_usuario = {}
for nombre_personalizado, nombre_variable in zip(nombres_personalizados, nombres_variables):
    label = tk.Label(root, text=nombre_personalizado + ":", padx=10, pady=5)
    label.grid(sticky="w", row=len(datos_usuario), column=0)
    entry = tk.Entry(root)
    entry.grid(row=len(datos_usuario), column=1)
    datos_usuario["{" + nombre_variable.replace(" ", "") + "}"] = entry

# Crear un frame para contener los botones
frame_botones = tk.Frame(root)
frame_botones.grid(row=len(datos_usuario) + 1, column=0, columnspan=3, padx=5, pady=10)

# Botón para seleccionar archivos
boton_seleccionar = tk.Button(frame_botones, text="Añadir documentos", command=seleccionar_archivos)
boton_seleccionar.grid(row=0, column=0, padx=5, pady=5)

# Botón para borrar campos
boton_borrar_campos = tk.Button(frame_botones, text="Borrar campos", command=borrar_campos)
boton_borrar_campos.grid(row=0, column=1, padx=5, pady=5)

# Botón para abrir la carpeta de archivos procesados
boton_abrir_carpeta = tk.Button(frame_botones, text="Procesados", command=abrir_carpeta_procesados)
boton_abrir_carpeta.grid(row=0, column=2, padx=5, pady=5)

# Botón para salir
boton_salir = tk.Button(frame_botones, text="Salir", command=salir)
boton_salir.grid(row=0, column=3, padx=5, pady=5)

# Configurar la alineación del frame
frame_botones.grid_columnconfigure(0, weight=1)
frame_botones.grid_columnconfigure(1, weight=1)
frame_botones.grid_columnconfigure(2, weight=1)
frame_botones.grid_columnconfigure(3, weight=1)

root.mainloop()